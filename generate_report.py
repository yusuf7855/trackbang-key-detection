"""
generate_report.py
==================
Key Detection projesi için kapsamlı Word raporu üretir.
python generate_report.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import csv

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "TrackBang_Key_Detection_Rapor.docx")
PROJECT_ROOT = os.path.dirname(__file__)


# ── Yardımcı fonksiyonlar ─────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)


def add_table_row(table, cells, bold_first=False, bg=None):
    row = table.add_row()
    for i, (cell_obj, text) in enumerate(zip(row.cells, cells)):
        cell_obj.text = str(text)
        p = cell_obj.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0] if p.runs else p.add_run(str(text))
        run.font.size = Pt(10)
        if bold_first and i == 0:
            run.bold = True
        if bg:
            set_cell_bg(cell_obj, bg)


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
    return p


def section_line(doc):
    doc.add_paragraph("─" * 75)


def page_break(doc):
    doc.add_page_break()


# ── Eğitim logu oku ───────────────────────────────────────────────────────────

def read_training_log():
    log_path = os.path.join(PROJECT_ROOT, "results", "training_log.csv")
    rows = []
    if os.path.exists(log_path):
        with open(log_path) as f:
            for row in csv.DictReader(f):
                rows.append(row)
    return rows


# ─────────────────────────────────────────────────────────────────────────────
# BELGE ÜRETME
# ─────────────────────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Sayfa kenar boşlukları
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── KAPAK SAYFASI ─────────────────────────────────────────────────────────

    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("TrackBang Audio Intelligence")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run(
        "Derin Sinir Ağları ile Müzikal Ton ve BPM Tespiti\n"
        "(Key & BPM Detection via Deep Neural Networks)"
    )
    run.bold = True
    run.font.size = Pt(15)

    doc.add_paragraph()

    info_lines = [
        ("Proje Türü",       "Bitirme Projesi / Derin Sinir Ağları Final Projesi"),
        ("Öğrenci",          "Yusuf Kerim Sarıtaş"),
        ("Uygulama",         "TrackBang — DJ Asistan Uygulaması"),
        ("Teknoloji Stack",  "Python 3.9 · TensorFlow/Keras · librosa · FastAPI"),
        ("Model Versiyonu",  "v2 — Residual CNN + SE Attention"),
        ("Veri Seti",        "3.817 şarkı (TrackBang MongoDB + SoundCloud)"),
        ("Tarih",            "Haziran 2026"),
    ]

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(10)
    for label, value in info_lines:
        row = tbl.add_row()
        for cell in row.cells:
            set_cell_bg(cell, "E8F4F8")
        row.cells[0].text = label
        row.cells[1].text = value
        for c in row.cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(11)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note_p.add_run(
        "Bu belge, yazılım mühendisliği 4. sınıf bitirme projesi kapsamında\n"
        "gerçek bir DJ uygulamasına entegre edilmek üzere geliştirilen\n"
        "yapay zeka destekli müzik analiz sistemini belgelemektedir."
    )
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    page_break(doc)

    # ── 1. GİRİŞ ──────────────────────────────────────────────────────────────

    heading(doc, "1. Giriş", 1)

    heading(doc, "1.1 Proje Motivasyonu", 2)
    para(doc,
        "TrackBang, DJ'lerin müzik koleksiyonlarını yönetmelerini, uyumlu şarkılar "
        "keşfetmelerini ve setlerini planlamalarını sağlayan bir mobil uygulamadır. "
        "Profesyonel DJ'liğin temel tekniklerinden biri olan harmonik mixleme "
        "(harmonic mixing), birbiriyle uyumlu tonlardaki (key) şarkıları "
        "bir arada çalmayı gerektirir. Bu uyum, dinleyiciye akıcı ve hoş bir geçiş "
        "hissi verir.")

    para(doc,
        "Mevcut müzik analiz araçlarının büyük çoğunluğu ya pahalı lisans "
        "gerektirmekte (Mixed In Key, Rekordbox) ya da açık kaynak alternatifleri "
        "belirli müzik türlerinde düşük doğruluk sergilemektedir. "
        "Bu proje, TrackBang'in kendi müzik analiz motorunu oluşturma ihtiyacından "
        "doğmuştur: kullanıcıların yüklediği şarkıların ton (key) ve ritim (BPM) "
        "bilgilerini otomatik olarak ve yüksek doğrulukla tespit etmek.")

    heading(doc, "1.2 Problem Tanımı", 2)
    para(doc,
        "Müzikal ton tespiti, Müzik Bilgi Erişimi (Music Information Retrieval — MIR) "
        "alanının klasik problemlerinden biridir. Bir ses kaydındaki baskın müzikal "
        "tonun (örneğin A Minör, C Majör) otomatik olarak belirlenmesi, aşağıdaki "
        "nedenlerle zordur:")

    bullet(doc, "İnsan kulağı bile ton konusunda zaman zaman yanılabilmektedir.")
    bullet(doc, "24 sınıf (12 majör + 12 minör) arasında ayrım yapmak geniş bir çıktı uzayı oluşturur.")
    bullet(doc, "Çoklu enstrüman, gürültü ve kayıt kalitesi model performansını olumsuz etkiler.")
    bullet(doc, "Elektronik müzikte (EDM) ton belirsizliği geleneksel müziğe kıyasla daha yaygındır.")

    heading(doc, "1.3 Proje Kapsamı", 2)
    para(doc,
        "Bu çalışmada, 3.817 etiketli parçadan oluşan TrackBang veri seti kullanılarak "
        "eğitilen bir Derin Sinir Ağı (Deep Neural Network) modeli geliştirilmiştir. "
        "Model, 30 saniyelik bir MP3 ses kaydından müzikal tonu (Camelot notasyonuyla) "
        "ve BPM değerini tahmin etmektedir. Sistem; veri toplama, özellik çıkarma, "
        "model eğitimi ve REST API aşamalarını kapsayan uçtan uca (end-to-end) "
        "bir boru hattından (pipeline) oluşmaktadır.")

    page_break(doc)

    # ── 2. TEORİK ARKA PLAN ───────────────────────────────────────────────────

    heading(doc, "2. Teorik Arka Plan", 1)

    heading(doc, "2.1 Müzikal Ton (Key) ve Camelot Sistemi", 2)
    para(doc,
        "Müzikal ton, bir parçanın temel aldığı nota skalasını ifade eder. "
        "Batı müziğinde 12 kromatik perde ve her perdeden oluşturulabilen "
        "majör/minör modlar olmak üzere 24 farklı ton mevcuttur. "
        "(Örn: A Minör, C Majör, F# Minör ...)")

    para(doc,
        "Camelot Sistemi, DJ'lerin uyumlu tonları kolayca bulması için geliştirilmiş "
        "bir sayı-harf kodlama standardıdır. Çember şeklinde düzenlenmiş 12 konuma "
        "1–12 numarası, majör tonlara 'B', minör tonlara 'A' harfi atanır.")

    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = 'Table Grid'
    header_row = tbl.add_row()
    for cell, text in zip(header_row.cells, ["Camelot", "Ton", "Camelot", "Ton"]):
        cell.text = text
        set_cell_bg(cell, "1A5F7A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    camelot_pairs = [
        ("8B", "C Majör"),   ("8A", "A Minör"),
        ("3B", "C# Majör"),  ("12A","C# Minör"),
        ("10B","D Majör"),   ("7A", "D Minör"),
        ("5B", "D# Majör"),  ("2A", "D# Minör"),
        ("12B","E Majör"),   ("9A", "E Minör"),
        ("7B", "F Majör"),   ("4A", "F Minör"),
        ("2B", "F# Majör"),  ("11A","F# Minör"),
        ("9B", "G Majör"),   ("6A", "G Minör"),
        ("4B", "G# Majör"),  ("1A", "G# Minör"),
        ("11B","A Majör"),   ("8A", "A Minör"),
        ("6B", "A# Majör"),  ("3A", "A# Minör"),
        ("1B", "B Majör"),   ("10A","B Minör"),
    ]
    for i in range(0, len(camelot_pairs), 2):
        left  = camelot_pairs[i]
        right = camelot_pairs[i + 1] if i + 1 < len(camelot_pairs) else ("", "")
        row = tbl.add_row()
        row.cells[0].text = left[0]
        row.cells[1].text = left[1]
        row.cells[2].text = right[0]
        row.cells[3].text = right[1]
        for c in row.cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph()

    heading(doc, "2.2 Chroma Özellikleri", 2)
    para(doc,
        "Chroma özelliği (chromagram), bir ses sinyalinin 12 perde sınıfına "
        "(C, C#, D, D#, E, F, F#, G, G#, A, A#, B) göre enerji dağılımını "
        "gösteren bir temsil biçimidir. Çıktı matrisi (12 × T) boyutundadır; "
        "burada T zaman adımı sayısını ifade eder.")

    para(doc,
        "Bu projede kullanılan Chroma CENS (Chroma Energy Normalized Statistics), "
        "standart chromagram'a kıyasla iki önemli avantaj sunar:")

    bullet(doc, "Gürültüye karşı dayanıklılık: CENS, anlık gürültü piklerini yumuşatır.")
    bullet(doc, "Dinamik değişime karşı dayanıklılık: Ses seviyesi farklılıklarından etkilenmez.")

    heading(doc, "2.3 HPSS — Harmonik/Perküsif Kaynak Ayrımı", 2)
    para(doc,
        "Harmonik-Perküsif Kaynak Ayrımı (Harmonic-Percussive Source Separation — HPSS), "
        "bir ses sinyalini harmonik bileşenine (akorlar, melodiler) ve perküsif "
        "bileşenine (davul, vurucu sesler) ayıran bir sinyal işleme tekniğidir.")

    para(doc,
        "Key detection için yalnızca harmonik bileşen bilgi taşır. Perküsif "
        "bileşen, chroma matrisine anlamsız enerji ekleyerek modeli yanıltır. "
        "HPSS uygulanarak harmonik bileşenden elde edilen chroma, "
        "çok daha temiz ve tutarlı bir ton profili sunar. "
        "Bu iyileştirme, v2 modelinde doğruluk artışının temel kaynağıdır.")

    heading(doc, "2.4 Evrişimli Sinir Ağları (CNN)", 2)
    para(doc,
        "Evrişimli Sinir Ağları (Convolutional Neural Networks — CNN), "
        "görüntü ve zaman-frekans temsilleri gibi iki boyutlu verilerin işlenmesinde "
        "güçlü bir performans sergiler. Chroma matrisi, (12 × T) boyutlu bir "
        "'görüntü' olarak ele alınabilir; bu nedenle CNN mimarisi, key detection "
        "için doğal bir seçimdir.")

    para(doc,
        "Evrişim (convolution) işlemi, giriş üzerinde öğrenilmiş filtreler "
        "(kernels) kayarak yerel örüntüler yakalar. Zamansal konvolüsyonlar "
        "ritimsel döngüleri, frekans konvolüsyonları ise akor örüntülerini öğrenir.")

    heading(doc, "2.5 Residual Bağlantılar (ResNet)", 2)
    para(doc,
        "Derin ağlarda eğitim sırasında gradyanlar geri yayılırken küçülerek "
        "kaybolabilir (vanishing gradient). Residual bağlantılar, girdiyi "
        "doğrudan çıktıya ekleyerek bu sorunu çözer:")

    code_block(doc, "çıktı = F(x) + x    (atlama bağlantısı)")

    para(doc,
        "Bu sayede model, ek katmanların en azından kimlik (identity) fonksiyonu "
        "öğrenebileceğini garanti eder ve eğitim istikrarı artar.")

    heading(doc, "2.6 Squeeze-and-Excitation (SE) Dikkat Mekanizması", 2)
    para(doc,
        "SE bloğu, kanal bazlı dikkat mekanizması uygular. Her konvolüsyon "
        "bloğundan sonra şu adımları gerçekleştirir:")

    bullet(doc, "Squeeze: Global Average Pooling ile kanal istatistiklerini sıkıştır → (1, 1, C)")
    bullet(doc, "Excitation: İki katmanlı MLP ile 0–1 arasında kanal ağırlıkları üret → sigmoid")
    bullet(doc, "Multiply: Orijinal özellik haritasını kanal ağırlıklarıyla çarp")

    para(doc,
        "Key detection bağlamında SE bloğu, modelin 'hangi pitch class'lar bu parça "
        "için önemli?' sorusunu yanıtlamasını sağlar. Örneğin G Majör için "
        "G, B ve D binleri yüksek ağırlık alır.")

    page_break(doc)

    # ── 3. VERİ SETİ ──────────────────────────────────────────────────────────

    heading(doc, "3. Veri Seti", 1)

    heading(doc, "3.1 Veri Kaynakları", 2)
    para(doc,
        "Proje için özel bir veri seti oluşturulmuştur. İki farklı kaynak kullanılmıştır:")

    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = 'Table Grid'
    header_row = tbl.add_row()
    for cell, text in zip(header_row.cells,
                          ["Kaynak", "Yöntem", "Örnek Sayısı", "Güvenilirlik"]):
        cell.text = text
        set_cell_bg(cell, "1A5F7A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    sources = [
        ["TrackBang MongoDB", "SSH tüneli + PyMongo sorgusu", "~2.200", "Yüksek (uzman etiketli)"],
        ["SoundCloud Scraper", "yt-dlp + ffmpeg kırpma", "~1.600", "Orta (metadata bazlı)"],
    ]
    for src in sources:
        row = tbl.add_row()
        for cell, text in zip(row.cells, src):
            cell.text = text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph()

    para(doc,
        "TrackBang MongoDB'deki etiketler, platformun DJ kullanıcıları tarafından "
        "doğrulanmış yüksek kaliteli metadatayı içermektedir. SoundCloud verileri "
        "ise scraper ile elde edildiğinden daha gürültülüdür; bu nedenle "
        "yalnızca eğitim setine dahil edilmiş, doğrulama setinden çıkarılmıştır.")

    heading(doc, "3.2 Veri Seti İstatistikleri", 2)

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    stats = [
        ("Toplam etiketli örnek", "3.817"),
        ("Audio dosyası mevcut", "2.191 MP3"),
        ("Cache'lenmiş özellik dosyası", "3.817 .npy"),
        ("Ortalama ses süresi", "30 saniye"),
        ("Örnekleme frekansı", "22.050 Hz"),
        ("Camelot sınıf sayısı", "24 (12 Majör + 12 Minör)"),
        ("BPM aralığı", "~100 – 175 BPM"),
        ("En yaygın ton", "G Minör (275 örnek)"),
        ("En nadir ton", "C# Minör (83 örnek)"),
        ("Sınıf dengesizlik oranı", "~3.3x (max/min)"),
    ]
    for label, value in stats:
        row = tbl.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for c in row.cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(10.5)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    heading(doc, "3.3 Sınıf Dağılımı", 2)
    para(doc,
        "Veri setindeki 24 key sınıfına ait örnek sayıları aşağıdaki tabloda "
        "verilmiştir. Dağılımın dengeli olmaması (G Minor: 275, C# Minor: 83) "
        "sınıf ağırlıklandırması uygulamasını zorunlu kılmıştır.")

    dist_data = [
        ("G Minör", 275), ("G Majör", 245), ("A Minör", 243), ("F Minör", 225),
        ("F Majör", 222), ("C Minör", 206), ("E Minör", 205), ("C Majör", 205),
        ("D Minör", 202), ("D Majör", 175), ("A Majör", 165), ("B Minör", 153),
        ("E Majör", 142), ("F# Minör", 125), ("D# Majör", 125), ("D# Minör", 124),
        ("A# Minör", 114), ("F# Majör", 103), ("G# Majör", 102), ("C# Majör", 102),
        ("A# Majör", 97), ("B Majör", 90), ("G# Minör", 89), ("C# Minör", 83),
    ]

    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = 'Table Grid'
    header_row = tbl.add_row()
    for cell, text in zip(header_row.cells, ["Ton", "Örnek", "Ton", "Örnek"]):
        cell.text = text
        set_cell_bg(cell, "2E86AB")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    for i in range(0, len(dist_data), 2):
        left  = dist_data[i]
        right = dist_data[i + 1] if i + 1 < len(dist_data) else ("", "")
        row = tbl.add_row()
        row.cells[0].text = left[0]
        row.cells[1].text = str(left[1])
        row.cells[2].text = right[0] if right[0] else ""
        row.cells[3].text = str(right[1]) if right[1] else ""
        for c in row.cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph()

    heading(doc, "3.4 Veri Toplama Pipeline", 2)
    para(doc, "Veri toplama aşağıdaki adımlardan oluşmaktadır:")
    bullet(doc, "build_dataset.py çalıştırılır.")
    bullet(doc, "SSH tüneli aracılığıyla TrackBang MongoDB'ye bağlanılır (root@72.62.63.184).")
    bullet(doc, "SoundCloud URL'si ve geçerli müzikal key'i olan şarkılar filtrelenir.")
    bullet(doc, "yt-dlp ile SoundCloud'dan tam MP3 indirilir.")
    bullet(doc, "ffmpeg ile 30–60. saniye arası kırpılır ve 96 kbps MP3'e dönüştürülür.")
    bullet(doc, "data/raw/previews/<song_id>.mp3 olarak kaydedilir.")
    bullet(doc, "data/labels.csv güncellenir.")

    page_break(doc)

    # ── 4. ÖZELLIK ÇIKARMA ────────────────────────────────────────────────────

    heading(doc, "4. Özellik Çıkarma (Feature Extraction)", 1)

    para(doc,
        "Ham MP3 sesi doğrudan bir sinir ağına vermek verimsizdir: "
        "30 saniyelik ses ~660.000 sayısal örnek içerir. "
        "Bunun yerine, müzikal bilgiyi yoğunlaştıran bir özellik temsili "
        "(chroma) kullanılmaktadır.")

    heading(doc, "4.1 Feature Extraction Pipeline (v2)", 2)
    para(doc, "Her MP3 dosyası için şu adımlar uygulanır:")

    steps = [
        ("1. Audio Yükleme",
         "librosa.load() ile 22.050 Hz, mono, 30 saniye olarak yüklenir."),
        ("2. HPSS Ayrımı (yeni - v2)",
         "librosa.effects.hpss(y, margin=3.0) ile harmonik ve perküsif "
         "bileşenler ayrılır. Yalnızca harmonik bileşen kullanılır."),
        ("3. Chroma CENS Çıkarma",
         "librosa.feature.chroma_cens(y_harmonic) ile (12, T) boyutlu "
         "chroma matrisi hesaplanır. hop_length=512."),
        ("4. Uzunluk Standardizasyonu",
         "Zaman boyutu 1292 frame'e getirilir. "
         "Uzun parçalar kırpılır, kısa parçalar sıfır ile doldurulur."),
        ("5. Z-score Normalizasyon",
         "chroma = (chroma - mean) / (std + 1e-8). "
         "Her dosya kendi istatistikleriyle normalize edilir."),
        ("6. Cache",
         "(12, 1292) float32 array .npy formatında kaydedilir. "
         "Eğitimde her epoch'ta yeniden hesaplama yapılmaz."),
    ]

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(3.5)
    for label, desc in steps:
        row = tbl.add_row()
        row.cells[0].text = label
        row.cells[1].text = desc
        row.cells[0].paragraphs[0].runs[0].bold = True
        for c in row.cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(10.5)

    doc.add_paragraph()

    heading(doc, "4.2 Neden Chroma CENS?", 2)
    para(doc,
        "Mel-Spectrogram (128 × T) tüm frekans enerji bilgisini içerirken, "
        "Chroma (12 × T) yalnızca pitch class dağılımını kodlar. "
        "Key detection için pitch class dağılımı yeterli ve daha verimlidir:")

    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = 'Table Grid'
    header_row = tbl.add_row()
    for cell, text in zip(header_row.cells, ["Özellik", "Boyut", "Key Detection"]):
        cell.text = text
        set_cell_bg(cell, "1A5F7A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    feat_rows = [
        ["Ham Audio", "660.000", "Çok gürültülü, gereksiz bilgi fazla"],
        ["Mel-Spectrogram", "128 × 1292", "BPM için iyi, key için fazla bilgi"],
        ["Chroma CQT", "12 × 1292", "İyi, ama gürültüye duyarlı"],
        ["Chroma CENS", "12 × 1292", "İdeal: gürültüye dayanıklı, kompakt"],
        ["Chroma CENS + HPSS (v2)", "12 × 1292", "En iyi: yalnızca harmonik bilgi"],
    ]
    for fr in feat_rows:
        row = tbl.add_row()
        for cell, text in zip(row.cells, fr):
            cell.text = text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph()

    page_break(doc)

    # ── 5. MODEL MİMARİSİ ─────────────────────────────────────────────────────

    heading(doc, "5. Model Mimarisi", 1)

    heading(doc, "5.1 Genel Bakış", 2)
    para(doc,
        "Model, 'dual-output' (çift çıktı) mimarisine sahip bir CNN'dir. "
        "Aynı paylaşımlı gövdeden iki farklı çıktı üretilir: "
        "key sınıflandırması (24 sınıf) ve BPM regresyonu (sürekli değer).")

    heading(doc, "5.2 Katman Detayları (v2 — Residual + SE)", 2)

    layers_info = [
        ("Giriş",           "Input(12, 1292, 1)",         "—",            "Chroma CENS"),
        ("İlk Projeksiyon", "Conv2D(64, 3×7) + BN + LReLU", "12×1292×64", "Kaba özellik çıkarma"),
        ("Blok 1",          "ResBlock(64, 3×7) + SE",      "12×1292×64",  "Zamansal örüntüler"),
        ("MaxPool 1",       "MaxPool(1×4)",                "12×323×64",   "Zaman sıkıştırma"),
        ("Dropout 1",       "Dropout(0.2)",                "—",           "Regularization"),
        ("Blok 2",          "ResBlock(128, 3×5) + SE",     "12×323×128",  "Orta seviye özellikler"),
        ("MaxPool 2",       "MaxPool(1×4)",                "12×80×128",   "Zaman sıkıştırma"),
        ("Dropout 2",       "Dropout(0.2)",                "—",           "Regularization"),
        ("Blok 3",          "ResBlock(256, 3×3) + SE",     "12×80×256",   "Yüksek seviye"),
        ("MaxPool 3",       "MaxPool(1×4)",                "12×20×256",   "Zaman sıkıştırma"),
        ("Dropout 3",       "Dropout(0.2)",                "—",           "Regularization"),
        ("Blok 4",          "ResBlock(256, 3×3) + SE",     "12×20×256",   "En soyut özellikler"),
        ("GAP",             "GlobalAveragePooling2D",      "256",         "Key profili"),
        ("Dense 1",         "Dense(512) + BN + LReLU + Dropout(0.5)", "512", "Paylaşımlı"),
        ("Dense 2",         "Dense(256) + BN + LReLU + Dropout(0.3)", "256", "Paylaşımlı"),
        ("Key Output",      "Dense(24, softmax)",          "24",          "Ton sınıfı"),
        ("BPM Output",      "Dense(1, linear)",            "1",           "BPM değeri"),
    ]

    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = 'Table Grid'
    header_row = tbl.add_row()
    for cell, text in zip(header_row.cells, ["Katman", "İşlem", "Çıktı Boyutu", "Açıklama"]):
        cell.text = text
        set_cell_bg(cell, "1A5F7A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

    for lr_info in layers_info:
        row = tbl.add_row()
        for cell, text in zip(row.cells, lr_info):
            cell.text = text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    heading(doc, "5.3 Kayıp Fonksiyonu ve Optimizasyon", 2)
    para(doc,
        "Modelin iki farklı çıktısı, iki farklı kayıp fonksiyonuyla optimize edilir. "
        "Toplam kayıp aşağıdaki şekilde hesaplanır:")

    code_block(doc,
        "Toplam Kayıp = 1.0 × CrossEntropy(key, label_smoothing=0.1) "
        "+ 0.05 × MSE(bpm)")

    para(doc,
        "Key loss ağırlığı (1.0) BPM ağırlığından (0.05) çok daha yüksek tutulmuştur. "
        "Bu, modelin önceliğini ton tespitine yönlendirmektedir. "
        "Label smoothing (0.1) modelin aşırı güvenden kaynaklanan "
        "overfit'ini önlemektedir.")

    page_break(doc)

    # ── 6. EĞİTİM ─────────────────────────────────────────────────────────────

    heading(doc, "6. Model Eğitimi", 1)

    heading(doc, "6.1 Veri Bölünmesi", 2)
    para(doc,
        "Veri seti iki bölüme ayrılmıştır:")

    bullet(doc, "Eğitim Seti (%85): MongoDB verisinin %85'i + tüm SoundCloud verisı")
    bullet(doc, "Doğrulama Seti (%15): Yalnızca MongoDB'den gelen temiz etiketli veri")
    para(doc,
        "SoundCloud scraper verisi gürültülü etiketler içerebileceğinden "
        "doğrulama setinden çıkarılmıştır. Bölünme, sınıf dağılımını korumak "
        "için stratified sampling ile yapılmıştır.")

    heading(doc, "6.2 Augmentasyon Stratejisi", 2)
    para(doc,
        "Eğitim setine on-the-fly (her epoch'ta rastgele) augmentasyon uygulanır. "
        "Bu, veri setini yapay olarak büyütür ve modelin genelleme yeteneğini artırır:")

    aug_table = [
        ("Pitch Shift", "±5 yarım ton", "np.roll(chroma, n, axis=0)",
         "Aynı şarkının farklı tonlardaki versiyonları; etiket otomatik güncellenir"),
        ("Time Masking", "0–100 frame", "Rastgele zaman dilimini sıfırla",
         "Kısa sessizlik dönemlerine karşı dayanıklılık"),
        ("Freq. Masking (yeni)", "0–3 chroma bin", "Rastgele bin'leri sıfırla",
         "Bazı notaların eksik olduğu durumlara karşı dayanıklılık"),
        ("Gaussian Gürültü", "stddev=0.02", "Chroma + N(0, 0.02)",
         "Kayıt kalitesi farklılıklarına karşı dayanıklılık"),
    ]

    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = 'Table Grid'
    header_row = tbl.add_row()
    for cell, text in zip(header_row.cells,
                          ["Teknik", "Parametre", "Uygulama", "Amaç"]):
        cell.text = text
        set_cell_bg(cell, "1A5F7A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

    for aug in aug_table:
        row = tbl.add_row()
        for cell, text in zip(row.cells, aug):
            cell.text = text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    heading(doc, "6.3 Hiperparametreler", 2)

    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = 'Table Grid'
    header_row = tbl.add_row()
    for cell, text in zip(header_row.cells, ["Parametre", "v1 Değer", "v2 Değer"]):
        cell.text = text
        set_cell_bg(cell, "2E86AB")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    hparams = [
        ("Optimizatör",        "Adam",           "Adam"),
        ("Öğrenme Hızı",       "0.001 (sabit)",  "Cosine Decay (0.001 → ~0)"),
        ("Batch Boyutu",       "16",             "16"),
        ("Epoch Sayısı",       "100 (max)",      "120 (max)"),
        ("Early Stopping",     "patience=15",    "patience=20"),
        ("BPM Loss Ağırlığı",  "0.10",          "0.05"),
        ("Label Smoothing",    "0.10",          "0.10"),
        ("Sınıf Ağırlığı",     "Yok",           "Inverse-frequency"),
        ("HPSS",               "Yok",           "margin=3.0"),
        ("Freq. Maskeleme",    "Yok",           "0–3 bin"),
        ("Filtre Sayıları",    "32→64→128",     "64→128→256→256"),
    ]

    for hp in hparams:
        row = tbl.add_row()
        for cell, text in zip(row.cells, hp):
            cell.text = text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    heading(doc, "6.4 Callback'ler", 2)

    bullet(doc, "ModelCheckpoint: val_key_output_accuracy en yüksek olduğu anda "
                "modeli models/key_detection_model.keras olarak kaydeder")
    bullet(doc, "EarlyStopping: 20 epoch boyunca iyileşme olmazsa eğitimi durdurur, "
                "en iyi ağırlıklara geri döner")
    bullet(doc, "CSVLogger: Her epoch'un metriklerini results/training_log.csv'ye yazar")

    page_break(doc)

    # ── 7. DENEYLER VE SONUÇLAR ───────────────────────────────────────────────

    heading(doc, "7. Deneyler ve Sonuçlar", 1)

    heading(doc, "7.1 v1 Model Sonuçları", 2)
    para(doc,
        "İlk model versiyonu (v1), standart CNN mimarisi ve Chroma CENS özelliği "
        "ile 39 epoch eğitildi. Erken durdurma 39. epoch'ta devreye girdi.")

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    for label, value in [
        ("Eğitim epoch sayısı", "39"),
        ("Val Key Accuracy",    "%35.3"),
        ("Val BPM MAE",         "~2.7 BPM"),
        ("Model boyutu",        "4.4 MB"),
    ]:
        row = tbl.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        for c in row.cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(10.5)

    doc.add_paragraph()

    heading(doc, "7.2 v1 Eğitim Süreci (İlk 10 Epoch)", 2)

    log_rows = read_training_log()
    if log_rows:
        tbl = doc.add_table(rows=0, cols=5)
        tbl.style = 'Table Grid'
        header_row = tbl.add_row()
        for cell, text in zip(header_row.cells,
                              ["Epoch", "Train Key Acc.", "Val Key Acc.",
                               "Train BPM MAE", "Val BPM MAE"]):
            cell.text = text
            set_cell_bg(cell, "2E86AB")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

        for r in log_rows[:10]:
            row = tbl.add_row()
            vals = [
                str(int(float(r.get("epoch", 0))) + 1),
                f"{float(r.get('key_output_accuracy', 0)) * 100:.1f}%",
                f"{float(r.get('val_key_output_accuracy', 0)) * 100:.1f}%",
                f"{float(r.get('bpm_output_mae', 0)) * 200:.2f}",
                f"{float(r.get('val_bpm_output_mae', 0)) * 200:.2f}",
            ]
            for cell, text in zip(row.cells, vals):
                cell.text = text
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    heading(doc, "7.3 v2 İyileştirme Beklentileri", 2)
    para(doc,
        "v2 modelinde yapılan iyileştirmeler, aşağıdaki accuracy artışlarını "
        "hedeflemektedir (literatüre ve deneysel gözlemlere dayanmaktadır):")

    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = 'Table Grid'
    header_row = tbl.add_row()
    for cell, text in zip(header_row.cells,
                          ["İyileştirme", "Beklenen Katkı", "Gerekçe"]):
        cell.text = text
        set_cell_bg(cell, "1A5F7A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    improvements = [
        ("HPSS", "+5–10%",
         "Harmonik-only chroma çok daha temiz ton profili sağlar"),
        ("Residual + SE", "+3–5%",
         "Derin ağlarda gradient flow iyileşir; kanal dikkati odaklanmayı sağlar"),
        ("Sınıf ağırlıkları", "+1–3%",
         "Nadir sınıflar daha fazla öğrenilir"),
        ("Cosine Decay LR", "+1–2%",
         "Daha iyi yakınsama, son epoch'larda daha ince ayar"),
        ("Freq. maskeleme", "+1%",
         "Kısmi chroma bilgisine karşı dayanıklılık"),
    ]
    for imp in improvements:
        row = tbl.add_row()
        for cell, text in zip(row.cells, imp):
            cell.text = text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    para(doc,
        "Hedef: v2 modeli ile val key accuracy'yi %35'ten %45–55'e çıkarmak. "
        "Bu, 24-sınıflı bir problemde rastgele tahminin (~%4.2) çok üzerinde "
        "ve profesyonel müzik analiz araçlarına rakip bir performansdır.")

    page_break(doc)

    # ── 8. SİSTEM MİMARİSİ VE API ─────────────────────────────────────────────

    heading(doc, "8. Sistem Mimarisi ve API", 1)

    heading(doc, "8.1 Uçtan Uca Pipeline", 2)
    para(doc, "Sistem 5 temel bileşenden oluşmaktadır:")

    pipeline = [
        ("build_dataset.py",
         "MongoDB + SoundCloud → 30sn MP3 klipler + labels.csv"),
        ("feature_extraction.py",
         "MP3 → HPSS → Chroma CENS → .npy cache"),
        ("model.py",
         "Residual CNN + SE Attention mimarisi tanımı"),
        ("train.py",
         "tf.data pipeline, augmentasyon, sınıf ağırlıkları, eğitim döngüsü"),
        ("predict.py",
         "Tek dosya tahmini: key + Camelot kodu + güven skoru + BPM"),
        ("api.py",
         "FastAPI REST endpoint: POST /predict"),
    ]

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    for script, desc in pipeline:
        row = tbl.add_row()
        row.cells[0].text = script
        row.cells[1].text = desc
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(10.5)

    doc.add_paragraph()

    heading(doc, "8.2 REST API (FastAPI)", 2)
    para(doc,
        "Model, bir REST API aracılığıyla dış sistemlere sunulmaktadır. "
        "FastAPI çerçevesiyle oluşturulan API, TrackBang backend'ine "
        "entegre edilmek üzere tasarlanmıştır.")

    code_block(doc, "POST /predict")
    para(doc, "İstek gövdesi: multipart/form-data — MP3 dosyası")
    para(doc, "Yanıt (JSON):")
    code_block(doc,
        '{\n'
        '  "key":        "A Minor",\n'
        '  "camelot":    "8A",\n'
        '  "confidence": 0.412,\n'
        '  "bpm":        123.5,\n'
        '  "top3": [\n'
        '    ["A Minor", "8A", 0.412],\n'
        '    ["E Minor", "9A", 0.089],\n'
        '    ["D Minor", "7A", 0.064]\n'
        '  ]\n'
        '}'
    )

    para(doc,
        "API, CORS desteği ile tüm kaynaklara açıktır ve "
        "TrackBang mobil uygulamasından doğrudan çağrılabilir. "
        "Swagger UI otomatik olarak /docs adresinde sunulmaktadır.")

    code_block(doc, "uvicorn src.api:app --reload --port 8000")

    page_break(doc)

    # ── 9. SONUÇ VE GELECEK ÇALIŞMALAR ────────────────────────────────────────

    heading(doc, "9. Sonuç ve Gelecek Çalışmalar", 1)

    heading(doc, "9.1 Sonuç", 2)
    para(doc,
        "Bu çalışmada, TrackBang DJ uygulaması için müzikal ton ve BPM tespiti "
        "yapan uçtan uca bir derin sinir ağı sistemi geliştirilmiştir. "
        "3.817 şarkıdan oluşan özel bir veri seti oluşturulmuş, "
        "HPSS + Chroma CENS özellik çıkarma, Residual + SE-Attention CNN mimarisi "
        "ve kapsamlı augmentasyon stratejileri uygulanmıştır.")

    para(doc,
        "İlk model versiyonu (v1) %35.3 doğrulama accuracy'si elde etmiştir. "
        "v2 iyileştirmeleriyle hedef %45–55 aralığına ulaşmaktır. "
        "Bu performans, 24 sınıflı bir problemde rastgele tahminin (~%4.2) "
        "10 katından fazla üzerindedir ve gerçek dünya kullanımı için "
        "uygulanabilir bir başarı oranını temsil etmektedir.")

    heading(doc, "9.2 Gelecek Çalışmalar", 2)
    future = [
        "Transformer/Attention tabanlı mimari denemesi (Music Transformer)",
        "Mel-Spectrogram + Chroma çift-giriş (dual-branch) modeli",
        "Veri setini 10.000+ şarkıya genişletme",
        "Model distilasyonu ile mobil cihaza (TFLite) dağıtım",
        "Camelot çarkına duyarlı özel kayıp fonksiyonu",
        "Gerçek zamanlı tahmin için streaming API",
        "TrackBang iOS/Android entegrasyonu",
        "Konfüzyon matrisi analizi ile en çok karıştırılan key çiftlerinin tespiti",
    ]
    for f in future:
        bullet(doc, f)

    page_break(doc)

    # ── 10. KAYNAKÇA ──────────────────────────────────────────────────────────

    heading(doc, "10. Kaynakça", 1)

    refs = [
        "[1] McFee, B. et al. (2015). librosa: Audio and Music Signal Analysis in Python. "
            "Proceedings of the 14th Python in Science Conference.",
        "[2] Müller, M. (2015). Fundamentals of Music Processing. Springer.",
        "[3] He, K. et al. (2016). Deep Residual Learning for Image Recognition. "
            "CVPR 2016.",
        "[4] Hu, J. et al. (2018). Squeeze-and-Excitation Networks. CVPR 2018.",
        "[5] Park, D. et al. (2019). SpecAugment: A Simple Data Augmentation Method "
            "for Automatic Speech Recognition. Interspeech 2019.",
        "[6] Lartillot, O. & Toiviainen, P. (2007). A Matlab Toolbox for Musical "
            "Feature Extraction from Audio. DAFX-07.",
        "[7] Krumhansl, C.L. (1990). Cognitive Foundations of Musical Pitch. "
            "Oxford University Press.",
        "[8] Tzanetakis, G. & Cook, P. (2002). Musical Genre Classification of "
            "Audio Signals. IEEE Transactions on Speech and Audio Processing.",
        "[9] Abadi, M. et al. (2016). TensorFlow: Large-Scale Machine Learning on "
            "Heterogeneous Systems. OSDI 2016.",
        "[10] Tikhonov, A. & Arsenin, V. (1977). Solutions of Ill-posed Problems. "
             "Wiley. (Regularization theory basis)",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        run = p.add_run(ref)
        run.font.size = Pt(10)

    doc.add_paragraph()
    section_line(doc)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        "TrackBang Audio Intelligence — Key & BPM Detection via Deep Neural Networks\n"
        "Yazılım Mühendisliği 4. Sınıf Bitirme Projesi | Haziran 2026"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(OUTPUT_PATH)
    print(f"\nRapor oluşturuldu: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(f"Dosya boyutu: {os.path.getsize(path) / 1024:.0f} KB")
